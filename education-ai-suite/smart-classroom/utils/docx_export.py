"""
Markdown to DOCX converter for class reports.

Converts the LLM-generated markdown report into a formatted Word document.
Handles: headings (h1-h3), paragraphs, bullet lists, numbered lists, bold, italic, tables.
"""

import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def markdown_to_docx(markdown_text: str, output_path: str, title: str = None) -> str:
    """Convert markdown text to a .docx file.

    Args:
        markdown_text: The markdown content to convert.
        output_path: Path where the .docx file will be saved.
        title: Optional document title (used as the first heading).

    Returns:
        The output_path on success.
    """
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Table detection
        if '|' in line and not in_table:
            # Check if this is a table (has separator row)
            if i + 1 < len(lines) and re.match(r'\s*\|[\s\-:|]+\|\s*$', lines[i + 1]):
                in_table = True
                table_rows = [line]
                i += 1
                continue

        if in_table:
            if '|' in line:
                # Skip separator row
                if not re.match(r'\s*\|[\s\-:|]+\|\s*$', line):
                    table_rows.append(line)
                i += 1
                continue
            else:
                # End of table
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
                # Don't increment i — process current line normally

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        # Bullet list
        elif re.match(r'^[\s]*[-*]\s', line):
            text = re.sub(r'^[\s]*[-*]\s', '', line)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text.strip())
        # Numbered list
        elif re.match(r'^[\s]*\d+\.\s', line):
            text = re.sub(r'^[\s]*\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text.strip())
        # Horizontal rule
        elif re.match(r'^---+\s*$', line):
            doc.add_paragraph('─' * 40)
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line.strip())

        i += 1

    # Flush remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"DOCX report saved to {output_path}")
    return output_path


def _add_formatted_text(paragraph, text: str):
    """Add text with bold and italic formatting to a paragraph."""
    # Split by bold (**text**) and italic (*text*) patterns
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _add_table(doc, rows: list):
    """Add a markdown table to the document."""
    if not rows:
        return

    def parse_row(row_text: str) -> list:
        cells = row_text.strip().strip('|').split('|')
        return [c.strip() for c in cells]

    header_cells = parse_row(rows[0])
    num_cols = len(header_cells)

    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'

    # Header row
    for j, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True

    # Data rows
    for row_text in rows[1:]:
        cells = parse_row(row_text)
        row = table.add_row()
        for j, cell_text in enumerate(cells[:num_cols]):
            row.cells[j].text = cell_text
