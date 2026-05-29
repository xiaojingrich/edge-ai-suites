"""
Report Template Manager.

Handles loading .docx templates, reading their text content for LLM,
and filling templates by replacing text in a copy (preserving formatting).

Template format:
- Any .docx file with placeholder text (XXX, XX, etc.)
- No special markup required — LLM decides what to replace
- Original template is never modified (always works on a copy)
"""

import os
import re
import json
import logging
import shutil
from pathlib import Path
from docx import Document

from utils.runtime_config_loader import RuntimeConfig

logger = logging.getLogger(__name__)

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")


def get_template_path(language: str = "zh", session_id: str = None) -> str:
    """Get the active template path. Checks for custom template first, then default."""
    if session_id:
        project_config = RuntimeConfig.get_section("Project")
        custom_path = os.path.join(
            project_config.get("location"),
            project_config.get("name"),
            session_id,
            "custom_report_template.docx",
        )
        if os.path.exists(custom_path):
            return custom_path

    project_config = RuntimeConfig.get_section("Project")
    project_custom = os.path.join(
        project_config.get("location"),
        project_config.get("name"),
        "report_template.docx",
    )
    if os.path.exists(project_custom):
        return project_custom

    default_path = os.path.join(TEMPLATES_DIR, f"report_template_{language}.docx")
    if os.path.exists(default_path):
        return default_path

    return None


def read_template_text(template_path: str) -> str:
    """Read all text content from a .docx template as plain text."""
    doc = Document(template_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                lines.append(" | ".join(row_cells))
    return "\n".join(lines)


def fill_template_from_text(template_path: str, replacements: dict, output_path: str) -> str:
    """Copy template and apply text replacements, preserving all formatting.

    Args:
        template_path: Source .docx template (never modified)
        replacements: Dict of {original_text: replacement_text}
        output_path: Where to save the filled copy

    Returns the output_path.
    """
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)

    doc = Document(output_path)

    for para in doc.paragraphs:
        _apply_replacements(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply_replacements(para, replacements)

    doc.save(output_path)
    logger.info(f"Template-based report saved to {output_path} ({len(replacements)} replacements)")
    return output_path


def _apply_replacements(paragraph, replacements: dict):
    """Apply text replacements in a paragraph while preserving run formatting."""
    full_text = paragraph.text
    if not full_text.strip():
        return

    new_text = full_text
    for original, replacement in replacements.items():
        if original in new_text:
            new_text = new_text.replace(original, replacement)

    if new_text == full_text:
        return

    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""


def parse_replacements_from_llm(response: str) -> dict:
    """Parse LLM response containing 'original → replacement' lines into a dict.

    Supports formats:
      - "原文 → 替换后"
      - "原文 -> 替换后"
    """
    replacements = {}
    for line in response.strip().split("\n"):
        line = line.strip()
        if not line:
            continue

        sep = None
        if "→" in line:
            sep = "→"
        elif "->" in line:
            sep = "->"
        else:
            continue

        parts = line.split(sep, 1)
        if len(parts) == 2:
            original = parts[0].strip().strip('"').strip("'")
            replacement = parts[1].strip().strip('"').strip("'")
            if original and replacement and original != replacement:
                replacements[original] = replacement

    return replacements
